"""DOCX + LaTeX Processing"""
import io, re
from typing import Tuple
from docx import Document

def extract_docx(data: bytes) -> Tuple[str, dict]:
    doc = Document(io.BytesIO(data))
    paras = [{"text": p.text.strip(), "style": p.style.name if p.style else "Normal"}
             for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(p["text"] for p in paras), {"paragraphs": paras}

def create_docx(text: str, meta: dict = None) -> bytes:
    doc = Document()
    for p in text.split("\n\n"):
        if p.strip(): doc.add_paragraph(p.strip())
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

def extract_latex(text: str) -> str:
    m = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', text, re.DOTALL)
    return m.group(1).strip() if m else text.strip()
